#!/usr/bin/env python3
"""
ACARA Curriculum Glossary Parser v1.0
=====================================

Parses ACARA Australian Curriculum V9.0 glossary documents (Word format) and 
converts them to CSV format suitable for loading into Neo4j or other graph databases.

GENERIC PARSER - Auto-detects subject from document.

ALIGNED WITH:
- Victorian Glossary Parser v3.0 (structure and CSV format)
- ACARA Curriculum Parser v4.2 (AU English conversion, text cleaning)

CSV OUTPUT FORMAT:
    SubjectArea,Subject,Band,Term,Definition

FEATURES:
- Auto-detects subject from document title (e.g., "English Glossary" → English)
- Extracts glossary terms from ACARA Word documents
- Handles tab-separated term/definition format in table cells
- Skips letter headers (A, B, C, etc.)
- Band-based organisation: Foundation to Year 10
- AU English spelling conversion
- Handles multi-line definitions

ACARA GLOSSARY FORMAT:
- Multiple tables containing glossary entries
- Single-column tables with term\tdefinition format
- Letter headers (A, B, C...) as section dividers
- First row typically contains "{Subject} Glossary"
- Terms may have trailing whitespace
- Definitions may span multiple lines

VERSION HISTORY:
v1.0 | 2026-01-31 | Initial release - generic parser with auto-detection

Requirements:
- Python 3.6+
- python-docx

Usage:
    python acara_glossary_parser_v1_0.py <input_file.docx> [output_dir]
    
Examples:
    python acara_glossary_parser_v1_0.py english-glossary-v9.docx
    python acara_glossary_parser_v1_0.py science-glossary-v9.docx /home/user/output

Author: Created for ACARA Curriculum V9.0 GraphRAG Project
License: MIT
"""

import re
import csv
import os
import sys
from pathlib import Path
from docx import Document
from typing import List, Dict, Optional, Tuple


# =============================================================================
# VERSION INFO
# =============================================================================

PARSER_VERSION = "1.0.0"
PARSER_DATE = "2026-01-31"


# =============================================================================
# US TO AU ENGLISH CONVERSION (ALIGNED WITH CURRICULUM PARSERS)
# =============================================================================

US_TO_AU_SPELLING = {
    # -ize to -ise (and derivatives)
    'recognize': 'recognise', 'recognizes': 'recognises', 'recognized': 'recognised', 'recognizing': 'recognising',
    'organize': 'organise', 'organizes': 'organises', 'organized': 'organised', 'organizing': 'organising',
    'analyze': 'analyse', 'analyzes': 'analyses', 'analyzed': 'analysed', 'analyzing': 'analysing',
    'summarize': 'summarise', 'summarizes': 'summarises', 'summarized': 'summarised', 'summarizing': 'summarising',
    'categorize': 'categorise', 'categorizes': 'categorises', 'categorized': 'categorised', 'categorizing': 'categorising',
    'prioritize': 'prioritise', 'prioritizes': 'prioritises', 'prioritized': 'prioritised', 'prioritizing': 'prioritising',
    'utilize': 'utilise', 'utilizes': 'utilises', 'utilized': 'utilised', 'utilizing': 'utilising',
    'minimize': 'minimise', 'minimizes': 'minimises', 'minimized': 'minimised', 'minimizing': 'minimising',
    'maximize': 'maximise', 'maximizes': 'maximises', 'maximized': 'maximised', 'maximizing': 'maximising',
    'emphasize': 'emphasise', 'emphasizes': 'emphasises', 'emphasized': 'emphasised', 'emphasizing': 'emphasising',
    'visualize': 'visualise', 'visualizes': 'visualises', 'visualized': 'visualised', 'visualizing': 'visualising',
    'realize': 'realise', 'realizes': 'realises', 'realized': 'realised', 'realizing': 'realising',
    'normalize': 'normalise', 'normalizes': 'normalises', 'normalized': 'normalised', 'normalizing': 'normalising',
    'standardize': 'standardise', 'standardizes': 'standardises', 'standardized': 'standardised', 'standardizing': 'standardising',
    'customize': 'customise', 'customizes': 'customises', 'customized': 'customised', 'customizing': 'customising',
    'specialize': 'specialise', 'specializes': 'specialises', 'specialized': 'specialised', 'specializing': 'specialising',
    'generalize': 'generalise', 'generalizes': 'generalises', 'generalized': 'generalised', 'generalizing': 'generalising',
    'memorize': 'memorise', 'memorizes': 'memorises', 'memorized': 'memorised', 'memorizing': 'memorising',
    'theorize': 'theorise', 'theorizes': 'theorises', 'theorized': 'theorised', 'theorizing': 'theorising',
    'hypothesize': 'hypothesise', 'hypothesizes': 'hypothesises', 'hypothesized': 'hypothesised', 'hypothesizing': 'hypothesising',
    'synthesize': 'synthesise', 'synthesizes': 'synthesises', 'synthesized': 'synthesised', 'synthesizing': 'synthesising',
    'criticize': 'criticise', 'criticizes': 'criticises', 'criticized': 'criticised', 'criticizing': 'criticising',
    'apologize': 'apologise', 'apologizes': 'apologises', 'apologized': 'apologised', 'apologizing': 'apologising',
    # -or to -our
    'color': 'colour', 'colors': 'colours', 'colored': 'coloured', 'coloring': 'colouring',
    'behavior': 'behaviour', 'behaviors': 'behaviours',
    'favor': 'favour', 'favors': 'favours', 'favored': 'favoured', 'favoring': 'favouring', 'favorable': 'favourable',
    'honor': 'honour', 'honors': 'honours', 'honored': 'honoured', 'honoring': 'honouring',
    'labor': 'labour', 'labors': 'labours', 'labored': 'laboured', 'laboring': 'labouring',
    'neighbor': 'neighbour', 'neighbors': 'neighbours', 'neighboring': 'neighbouring',
    'humor': 'humour', 'humors': 'humours',
    'vigor': 'vigour',
    'endeavor': 'endeavour', 'endeavors': 'endeavours',
    # -er to -re
    'center': 'centre', 'centers': 'centres', 'centered': 'centred', 'centering': 'centring',
    'meter': 'metre', 'meters': 'metres',
    'theater': 'theatre', 'theaters': 'theatres',
    'fiber': 'fibre', 'fibers': 'fibres',
    # -og to -ogue
    'catalog': 'catalogue', 'catalogs': 'catalogues',
    'dialog': 'dialogue', 'dialogs': 'dialogues',
    'analog': 'analogue',
    'prolog': 'prologue',
    'epilog': 'epilogue',
    # Other common differences
    'program': 'programme', 'programs': 'programmes',
    'defense': 'defence',
    'offense': 'offence',
    'license': 'licence',
    'practice': 'practise',
    'judgment': 'judgement',
    'aging': 'ageing',
    'modeling': 'modelling',
    'traveling': 'travelling', 'traveled': 'travelled', 'traveler': 'traveller',
    'canceled': 'cancelled', 'canceling': 'cancelling',
    'labeled': 'labelled', 'labeling': 'labelling',
    'leveled': 'levelled', 'leveling': 'levelling',
    'signaled': 'signalled', 'signaling': 'signalling',
}


def convert_to_au_english(text: str) -> str:
    """Convert US English spellings to Australian English.
    
    Processes text to ensure consistent AU English spelling in output.
    Case-insensitive matching with case preservation.
    """
    if not text:
        return text
    
    result = text
    for us_spelling, au_spelling in US_TO_AU_SPELLING.items():
        pattern = re.compile(re.escape(us_spelling), re.IGNORECASE)
        
        def replace_preserve_case(match):
            matched = match.group(0)
            if matched.isupper():
                return au_spelling.upper()
            elif matched[0].isupper():
                return au_spelling.capitalize()
            return au_spelling
        
        result = pattern.sub(replace_preserve_case, result)
    
    return result


# =============================================================================
# TEXT CLEANING
# =============================================================================

def capitalise_first(text: str) -> str:
    """Capitalise first letter, preserve rest."""
    if not text:
        return ""
    return text[0].upper() + text[1:] if len(text) > 1 else text.upper()


def clean_text(text: str) -> str:
    """Clean text by removing extra whitespace and converting to AU English."""
    if not text:
        return ""
    # Replace multiple spaces/tabs with single space
    text = re.sub(r'[^\S\n]+', ' ', text)
    # Clean up multiple newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)
    # Strip leading/trailing whitespace from each line
    lines = text.split('\n')
    lines = [line.strip() for line in lines]
    cleaned = '\n'.join(lines).strip()
    # Apply AU English spelling conversion
    return convert_to_au_english(cleaned)


def is_letter_header(text: str) -> bool:
    """Check if text is a single letter header (A, B, C, etc.)."""
    text = text.strip()
    return len(text) == 1 and text.isalpha() and text.isupper()


# =============================================================================
# SUBJECT DETECTION
# =============================================================================

# Known ACARA subjects for auto-detection
KNOWN_SUBJECTS = {
    'english': ('English', 'English'),
    'mathematics': ('Mathematics', 'Mathematics'),
    'science': ('Science', 'Science'),
    'hass': ('Humanities and Social Sciences', 'HASS'),
    'history': ('Humanities and Social Sciences', 'History'),
    'geography': ('Humanities and Social Sciences', 'Geography'),
    'civics and citizenship': ('Humanities and Social Sciences', 'Civics and Citizenship'),
    'economics and business': ('Humanities and Social Sciences', 'Economics and Business'),
    'health and physical education': ('Health and Physical Education', 'Health and Physical Education'),
    'hpe': ('Health and Physical Education', 'Health and Physical Education'),
    'the arts': ('The Arts', 'The Arts'),
    'visual arts': ('The Arts', 'Visual Arts'),
    'music': ('The Arts', 'Music'),
    'drama': ('The Arts', 'Drama'),
    'dance': ('The Arts', 'Dance'),
    'media arts': ('The Arts', 'Media Arts'),
    'technologies': ('Technologies', 'Technologies'),
    'digital technologies': ('Technologies', 'Digital Technologies'),
    'design and technologies': ('Technologies', 'Design and Technologies'),
    'languages': ('Languages', 'Languages'),
}


def detect_subject_from_document(doc: Document) -> Tuple[str, str]:
    """
    Auto-detect subject from document content.
    
    Looks for patterns like "{Subject} Glossary" in the document.
    
    Args:
        doc: Parsed Word document
        
    Returns:
        Tuple of (subject_area, subject)
    """
    # Search first few cells for glossary title
    for table in doc.tables:
        for row in table.rows[:5]:
            for cell in row.cells:
                text = cell.text.strip().lower()
                
                # Look for "{Subject} Glossary" pattern
                if 'glossary' in text:
                    # Extract subject name before "glossary"
                    match = re.match(r'(.+?)\s*glossary', text, re.IGNORECASE)
                    if match:
                        subject_text = match.group(1).strip().lower()
                        
                        # Look up in known subjects
                        if subject_text in KNOWN_SUBJECTS:
                            return KNOWN_SUBJECTS[subject_text]
                        
                        # Try partial matches
                        for key, (area, subj) in KNOWN_SUBJECTS.items():
                            if key in subject_text or subject_text in key:
                                return (area, subj)
                        
                        # Fallback: use detected text as both area and subject
                        detected = subject_text.title()
                        return (detected, detected)
    
    return (None, None)


def detect_subject_from_filename(filename: str) -> Tuple[str, str]:
    """
    Detect subject from filename as fallback.
    
    Args:
        filename: Input filename
        
    Returns:
        Tuple of (subject_area, subject)
    """
    name = Path(filename).stem.lower()
    
    for key, (area, subj) in KNOWN_SUBJECTS.items():
        if key.replace(' ', '-') in name or key.replace(' ', '_') in name or key.replace(' ', '') in name:
            return (area, subj)
    
    return (None, None)


# =============================================================================
# GLOSSARY PARSER CLASS
# =============================================================================

class ACARAGlossaryParser:
    """Generic parser for ACARA curriculum glossary documents."""
    
    def __init__(self, input_file: str, output_dir: Optional[str] = None):
        """
        Initialise parser with input file.
        
        Args:
            input_file: Path to glossary Word document
            output_dir: Optional output directory (defaults to same as input)
        """
        self.input_file = input_file
        self.output_dir = output_dir or os.path.dirname(input_file) or '.'
        self.band = 'Foundation to Year 10'  # F-10 glossaries
        self.glossary_terms: List[Dict] = []
        
        # Load document and detect subject
        self.doc = Document(input_file)
        self.subject_area, self.subject = self._detect_subject()
        
        os.makedirs(self.output_dir, exist_ok=True)
    
    def _detect_subject(self) -> Tuple[str, str]:
        """Detect subject from document content or filename."""
        # Try document content first
        subject_area, subject = detect_subject_from_document(self.doc)
        
        if subject:
            return (subject_area, subject)
        
        # Fallback to filename
        subject_area, subject = detect_subject_from_filename(self.input_file)
        
        if subject:
            return (subject_area, subject)
        
        # Final fallback: ask user or use Unknown
        print("⚠️  Could not auto-detect subject from document or filename.")
        print("    Please ensure the document contains '{Subject} Glossary' title")
        print("    or the filename contains the subject name.")
        return ('Unknown', 'Unknown')
    
    def parse_glossary(self) -> List[Dict]:
        """
        Parse glossary from ACARA Word document.
        
        ACARA glossary format:
        - Multiple tables with single-column cells
        - Each cell contains: term\tdefinition (tab-separated)
        - Letter headers (A, B, C...) as section dividers
        - Title row "{Subject} Glossary" to skip
        
        Returns:
            List of glossary term dictionaries
        """
        print(f"\nParsing: {os.path.basename(self.input_file)}")
        
        terms_found = 0
        skipped_headers = 0
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    
                    if not cell_text:
                        continue
                    
                    # Skip title rows (e.g., "English Glossary")
                    if 'glossary' in cell_text.lower() and len(cell_text) < 50:
                        continue
                    
                    # Skip letter headers
                    if is_letter_header(cell_text):
                        skipped_headers += 1
                        continue
                    
                    # Extract term and definition
                    # Format 1: tab-separated (most common): "term\tdefinition"
                    # Format 2: newline-separated (some entries): "term\ndefinition"
                    term = None
                    definition = None
                    
                    if '\t' in cell_text:
                        # Tab-separated: split on first tab
                        parts = cell_text.split('\t', 1)
                        term = parts[0].strip()
                        definition = parts[1].strip() if len(parts) > 1 else ""
                        # Clean up definition (may have leading newline)
                        definition = definition.lstrip('\n').strip()
                    elif '\n' in cell_text:
                        # Newline-separated: first line is term, rest is definition
                        lines = cell_text.split('\n', 1)
                        term = lines[0].strip()
                        definition = lines[1].strip() if len(lines) > 1 else ""
                    
                    if term and definition:
                        self.glossary_terms.append({
                            'SubjectArea': self.subject_area,
                            'Subject': self.subject,
                            'Band': self.band,
                            'Term': capitalise_first(clean_text(term)),
                            'Definition': clean_text(definition)
                        })
                        terms_found += 1
        
        print(f"  ✓ Found {terms_found} glossary terms")
        print(f"  ✓ Skipped {skipped_headers} letter headers")
        
        return self.glossary_terms
    
    def write_csv(self) -> str:
        """
        Write glossary terms to CSV file.
        
        CSV format: SubjectArea,Subject,Band,Term,Definition
        (Aligned with Victorian glossary and import patterns skill)
        
        Returns:
            Path to output file
        """
        filename = f"ACARA {self.subject} V9 - Glossary.csv"
        output_file = os.path.join(self.output_dir, filename)
        
        with open(output_file, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['SubjectArea', 'Subject', 'Band', 'Term', 'Definition'])
            
            # Sort by term alphabetically
            sorted_terms = sorted(self.glossary_terms, key=lambda x: x['Term'].lower())
            
            for term in sorted_terms:
                writer.writerow([
                    term['SubjectArea'],
                    term['Subject'],
                    term['Band'],
                    term['Term'],
                    term['Definition']
                ])
        
        print(f"\n{filename}: {len(self.glossary_terms)} terms")
        return output_file
    
    def validate(self) -> bool:
        """Validate parsed glossary data."""
        print("\n" + "=" * 60)
        print("VALIDATION")
        print("=" * 60)
        
        print(f"Total terms: {len(self.glossary_terms)}")
        
        # Check for duplicates
        terms = [t['Term'].lower() for t in self.glossary_terms]
        dups = [t for t in set(terms) if terms.count(t) > 1]
        if dups:
            print(f"⚠️  Duplicate terms: {dups}")
        
        # Check for empty definitions
        empty_defs = [t['Term'] for t in self.glossary_terms if not t['Definition']]
        if empty_defs:
            print(f"⚠️  Terms with empty definitions: {empty_defs}")
        
        # Show sample
        print("\nSample terms (first 5):")
        for term in self.glossary_terms[:5]:
            def_preview = term['Definition'][:60] + "..." if len(term['Definition']) > 60 else term['Definition']
            print(f"  - {term['Term']}: {def_preview}")
        
        return len(dups) == 0 and len(empty_defs) == 0
    
    def run(self) -> str:
        """Execute full parsing pipeline."""
        print("\n" + "=" * 70)
        print(f"ACARA GLOSSARY PARSER v{PARSER_VERSION}")
        print("=" * 70)
        print(f"Input:        {self.input_file}")
        print(f"Subject Area: {self.subject_area}")
        print(f"Subject:      {self.subject}")
        print(f"Band:         {self.band}")
        print(f"Output:       {self.output_dir}")
        print("")
        
        # Parse glossary
        self.parse_glossary()
        
        # Validate
        self.validate()
        
        # Write CSV
        print("\n" + "=" * 60)
        print("WRITING CSV")
        print("=" * 60)
        
        output_path = self.write_csv()
        
        print("\n" + "=" * 70)
        print("COMPLETE!")
        print(f"Output: {output_path}")
        print("=" * 70)
        
        return output_path


# =============================================================================
# MAIN EXECUTION
# =============================================================================

def main():
    """Main entry point with command-line support."""
    
    # Parse command line arguments
    if len(sys.argv) < 2:
        print(f"ACARA Glossary Parser v{PARSER_VERSION}")
        print("")
        print("Usage: python acara_glossary_parser_v1_0.py <input_file.docx> [output_dir]")
        print("")
        print("Arguments:")
        print("  input_file.docx  Path to ACARA glossary Word document")
        print("  output_dir       Optional output directory (default: same as input)")
        print("")
        print("Examples:")
        print("  python acara_glossary_parser_v1_0.py english-glossary-v9.docx")
        print("  python acara_glossary_parser_v1_0.py science-glossary-v9.docx ./output")
        print("")
        print("The parser auto-detects the subject from the document.")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    # Validate input file
    if not os.path.exists(input_file):
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)
    
    if not input_file.lower().endswith('.docx'):
        print(f"Error: Input file must be a .docx file: {input_file}")
        sys.exit(1)
    
    # Run parser
    parser = ACARAGlossaryParser(input_file, output_dir)
    parser.run()


if __name__ == '__main__':
    main()
