#!/usr/bin/env python3
"""
ACARA Curriculum Parser v4.10
=============================

A generic parser for ACARA Australian Curriculum V9.0 documents.

ALIGNED WITH VICTORIAN PARSER v4.2:
- Standardised keyword extraction (min word length 3, no hard limit)
- Standardised stop words list (merged best from both parsers)
- Generic strand inference using keyword matching (not subject-specific)
- Nested table handling via get_cell_text() function
- Three-phase test infrastructure (Pre-check, Parse, Post-check)
- AS Components CSV format aligned (SubjectArea, Subject, Band, ASComponentCode, ASComponentText, ...)
- Glossary CSV includes Band column

IMPORTANT - SubjectArea Consistency:
- All output files (Curriculum, AS Components, Glossary) use the SAME subject_area from config
- For HASS subjects (Geography, History, Civics, Economics), use 'HASS' as subject_area
- For The Arts subjects (Dance, Drama, Media Arts, Music, Visual Arts), use 'The Arts' as subject_area
- This ensures Neo4j loads create consistent SubjectArea nodes across all file types

VERSION HISTORY:
---------------
v4.10.0 | 2026-02-03 | FIX: The Arts F-6 band mapping corrected
                     | The Arts F-6 uses BANDED structure (Years 1 and 2, Years 3 and 4, Years 5 and 6)
                     | NOT single year bands as previously documented in v4.9.0
                     | map_as_year_to_band() now maps Year 2->Years 1 and 2, Year 4->Years 3 and 4, etc.
                     | FIX: Foundation Achievement Standard regex
                     | The Arts documents use "By the end of the Foundation year" (not "By the end of Foundation")
                     | Updated regex pattern to match both formats: (?:the )?(Foundation(?: year)?|Year \d+)
                     | NOTE: The Arts subject codes use 3-letter format (AMU, ADA, ADR, AMA, AVA)
v4.9.0 | 2026-02-03 | FIX: Achievement Standards extraction for The Arts 7-10 documents
                    | _parse_standard_achievement_standards() now checks for BOTH
                    | 'Year level description' AND 'Band level description' table headers
                    | The Arts subjects (Dance, Drama, Media Arts, Music, Visual Arts) use
                    | 'Band level description' in 7-10 documents, causing AS to be missed
                    | FIX: map_as_year_to_band() extended for The Arts subject area
                    | Year 8 -> Years 7 and 8, Year 10 -> Years 9 and 10
v4.8.0 | 2026-02-03 | FEATURE: Glossary term capitalisation
                    | Glossary terms now have first letter capitalised for consistency
                    | Uses existing capitalise_first() function
v4.7.0 | 2026-02-03 | FEATURE: Technologies subject support (Digital Technologies, Design and Technologies)
                    | Technologies subjects have dual AS tables: subject-specific and Learning area
                    | Parser now uses 'Learning area Achievement standard' for Technologies (more comprehensive)
                    | Added subject_area detection to route Technologies to correct AS parsing logic
                    | F-6: Uses Row 5 of description tables (Learning area AS)
                    | 7-10: Uses standalone 'Learning area Achievement standard' table, or fallback to Row 3
                    | FEATURE: Band name normalisation for VCAA consistency
                    | ACARA "Years X–Y" (en-dash) -> "Years X and Y" format
                    | Added normalise_acara_band() and map_as_year_to_band() functions
                    | FIX: Parser now loops through all curriculum_docs (not just first)
v4.6.0 | 2026-02-02 | FEATURE: Topic extraction for History subjects
                    | Detects embedded topics (Greece, Rome, Medieval Europe, etc.) in elaboration cells
                    | Topics are short lines (<60 chars) not starting with elaboration verbs
                    | Each elaboration paired with its topic in CSV output
                    | Works for both ACARA and Victorian History documents
v4.5.0 | 2026-02-02 | DOC: Clarified SubjectArea consistency requirement for HASS subjects
                    | All outputs now clearly documented to use same subject_area from config
v4.4.0 | 2026-02-02 | FIX: Glossary terms with / in name (e.g., "Asia / Asian") now parsed correctly
                    | These terms use newline separator instead of tab; parser now handles both formats
v4.3.0 | 2026-02-01 | FIX: Duplicate elaborations caused by capitalise_first() applied after duplicate check
v4.2.0 | 2026-01-30 | ALIGNMENT with Victorian parser v4.2 - CSV column order standardised
v4.1.0 | 2026-01-30 | Added ASComponentCode generation (AC9E{band}ASC{seq})
v4.0.0 | 2026-01-30 | ALIGNMENT with Victorian parser v4.0
v3.2.0 | 2026-01-15 | Integrated comprehensive test suite
v3.1.1 | 2026-01-15 | Added capitalisation
v3.1.0 | 2026-01-15 | Added EALD support
v3.0.0 | 2026-01-13 | AS Components with confidence scoring
"""

import csv
import re
import os
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple, Set
from docx import Document
from collections import defaultdict


# =============================================================================
# VERSION INFO
# =============================================================================

PARSER_VERSION = "4.10.0"
PARSER_DATE = "2026-02-03"


# =============================================================================
# CONFIGURATION - CHANGE THESE FOR YOUR SUBJECT
# =============================================================================

# IMPORTANT: subject_area value is used in ALL output files (Curriculum, AS Components, Glossary)
# For HASS subjects, use 'HASS' (not 'Humanities and Social Sciences') to match Neo4j database

SUBJECT_CONFIG = {
    'subject_area': 'The Arts',
    'subject': 'Music',
    'subject_code': 'AMU',
    'bands_f6': ['Foundation', 'Years 1 and 2', 'Years 3 and 4', 'Years 5 and 6'],
    'bands_7_10': ['Years 7 and 8', 'Years 9 and 10'],
    'strands': None,  # Auto-discovered from document
    'strand_codes': {  # Music strand codes - to be confirmed from document
        'E': 'Exploring and responding',
        'D': 'Developing practices and skills',
        'C': 'Creating and making'
    },
    'curriculum_docs': [
        '/mnt/user-data/uploads/the-arts-music-curriculum-content-f-6-v9.docx',
        '/mnt/user-data/uploads/the-arts-music-curriculum-content-7-10-v9.docx'
    ],
    'eald_docs': [],
    'glossary_doc': '/mnt/user-data/uploads/the-arts-music-glossary-v9.docx',
    'output_dir': '/home/claude/acara_music_output'
}

# =============================================================================
# EXAMPLE CONFIGURATIONS FOR OTHER SUBJECTS
# =============================================================================
#
# GEOGRAPHY (HASS - Years 7-10 only):
# SUBJECT_CONFIG = {
#     'subject_area': 'HASS',  # CRITICAL: Use 'HASS' not 'Humanities and Social Sciences'
#     'subject': 'Geography',
#     'subject_code': 'HG',
#     'bands_f6': [],  # No F-6 for Geography
#     'bands_7_10': ['Year 7', 'Year 8', 'Year 9', 'Year 10'],
#     'strand_codes': {'K': 'Knowledge and understanding', 'S': 'Skills'},
#     'curriculum_docs': ['path/to/geography-curriculum-content-v9.docx'],
#     'eald_docs': [],  # Optional for non-English
#     'glossary_doc': 'path/to/geography-glossary-v9.docx',
#     'output_dir': '/home/claude/acara_geography_output'
# }
#
# HISTORY (HASS - Years 7-10 only):
# NOTE: History has embedded topics (Greece, Rome, Medieval Europe, etc.) that are
# automatically extracted. Topics appear in the Topic column of the CSV output.
# SUBJECT_CONFIG = {
#     'subject_area': 'HASS',  # CRITICAL: Use 'HASS' not 'Humanities and Social Sciences'
#     'subject': 'History',    # CRITICAL: Must be 'History' for topic extraction to work
#     'subject_code': 'HH',
#     'bands_f6': [],
#     'bands_7_10': ['Year 7', 'Year 8', 'Year 9', 'Year 10'],
#     'strand_codes': {'K': 'Knowledge and understanding', 'S': 'Skills'},
#     'curriculum_docs': ['path/to/history-curriculum-content-v9.docx'],
#     'eald_docs': [],
#     'glossary_doc': 'path/to/history-glossary-v9.docx',
#     'output_dir': '/home/claude/acara_history_output'
# }
#
# CIVICS AND CITIZENSHIP (HASS - Years 7-10 only):
# SUBJECT_CONFIG = {
#     'subject_area': 'HASS',
#     'subject': 'Civics and Citizenship',
#     'subject_code': 'HC',
#     ...
# }
#
# ECONOMICS AND BUSINESS (HASS - Years 7-10 only):
# SUBJECT_CONFIG = {
#     'subject_area': 'HASS',
#     'subject': 'Economics and Business',
#     'subject_code': 'HE',
#     ...
# }
#
# =============================================================================
# TECHNOLOGIES SUBJECTS - Special AS handling
# =============================================================================
# Technologies subjects have TWO achievement standards in documents:
# - Subject-specific AS (shorter)
# - Learning area Achievement standard (comprehensive - THIS IS USED)
#
# DIGITAL TECHNOLOGIES (Foundation to Year 10):
# SUBJECT_CONFIG = {
#     'subject_area': 'Technologies',  # CRITICAL: Use 'Technologies' for correct AS parsing
#     'subject': 'Digital Technologies',
#     'subject_code': 'TDI',  # Foundation uses TDIF pattern
#     'bands_f6': ['Foundation', 'Years 1 and 2', 'Years 3 and 4', 'Years 5 and 6'],
#     'bands_7_10': ['Years 7 and 8', 'Years 9 and 10'],
#     'strand_codes': {'K': 'Knowledge and understanding', 'P': 'Processes and production skills'},
#     'curriculum_docs': [
#         'path/to/technologies-digital-technologies-curriculum-content-f-6-v9.docx',
#         'path/to/technologies-digital-technologies-curriculum-content-7-10-v9.docx'
#     ],
#     'eald_docs': [],
#     'glossary_doc': 'path/to/technologies-digital-technologies-glossary-v9.docx',
#     'output_dir': '/home/claude/acara_digitech_output'
# }
#
# DESIGN AND TECHNOLOGIES (Foundation to Year 10):
# SUBJECT_CONFIG = {
#     'subject_area': 'Technologies',  # CRITICAL: Use 'Technologies' for correct AS parsing
#     'subject': 'Design and Technologies',
#     'subject_code': 'TDE',  # Foundation uses TDEF pattern
#     'bands_f6': ['Foundation', 'Years 1 and 2', 'Years 3 and 4', 'Years 5 and 6'],
#     'bands_7_10': ['Years 7 and 8', 'Years 9 and 10'],
#     'strand_codes': {'K': 'Knowledge and understanding', 'P': 'Processes and production skills'},
#     'curriculum_docs': [
#         'path/to/technologies-design-and-technologies-curriculum-content-f-6-v9.docx',
#         'path/to/technologies-design-and-technologies-curriculum-content-7-10-v9.docx'
#     ],
#     'eald_docs': [],
#     'glossary_doc': 'path/to/technologies-design-and-technologies-glossary-v9.docx',
#     'output_dir': '/home/claude/acara_designtech_output'
# }

# =============================================================================
# THE ARTS SUBJECTS - F-10 with banded structure
# =============================================================================
# The Arts subjects use 3-letter subject codes: AMU, ADA, ADR, AMA, AVA
# All use banded structure for F-6: Foundation, Years 1 and 2, Years 3 and 4, Years 5 and 6
# All use banded structure for 7-10: Years 7 and 8, Years 9 and 10
# Documents use "By the end of the Foundation year" (not "By the end of Foundation")
#
# MUSIC (Foundation to Year 10):
# SUBJECT_CONFIG = {
#     'subject_area': 'The Arts',  # CRITICAL: Use 'The Arts' for correct AS band mapping
#     'subject': 'Music',
#     'subject_code': 'AMU',  # Arts Music
#     'bands_f6': ['Foundation', 'Years 1 and 2', 'Years 3 and 4', 'Years 5 and 6'],
#     'bands_7_10': ['Years 7 and 8', 'Years 9 and 10'],
#     'strand_codes': {
#         'E': 'Exploring and responding',
#         'D': 'Developing practices and skills',
#         'C': 'Creating and making',
#         'P': 'Presenting and performing'
#     },
#     'curriculum_docs': [
#         'path/to/the-arts-music-curriculum-content-f-6-v9.docx',
#         'path/to/the-arts-music-curriculum-content-7-10-v9.docx'
#     ],
#     'eald_docs': [],
#     'glossary_doc': 'path/to/the-arts-music-glossary-v9.docx',
#     'output_dir': '/home/claude/acara_music_output'
# }
#
# DANCE (Foundation to Year 10):
# SUBJECT_CONFIG = {
#     'subject_area': 'The Arts',
#     'subject': 'Dance',
#     'subject_code': 'ADA',  # Arts Dance
#     'bands_f6': ['Foundation', 'Years 1 and 2', 'Years 3 and 4', 'Years 5 and 6'],
#     'bands_7_10': ['Years 7 and 8', 'Years 9 and 10'],
#     ...
# }
#
# DRAMA (Foundation to Year 10):
# SUBJECT_CONFIG = {
#     'subject_area': 'The Arts',
#     'subject': 'Drama',
#     'subject_code': 'ADR',  # Arts Drama
#     ...
# }
#
# MEDIA ARTS (Foundation to Year 10):
# SUBJECT_CONFIG = {
#     'subject_area': 'The Arts',
#     'subject': 'Media Arts',
#     'subject_code': 'AMA',  # Arts Media Arts
#     ...
# }
#
# VISUAL ARTS (Foundation to Year 10):
# SUBJECT_CONFIG = {
#     'subject_area': 'The Arts',
#     'subject': 'Visual Arts',
#     'subject_code': 'AVA',  # Arts Visual Arts
#     ...
# }


def get_acara_code_pattern(subject_code: str) -> re.Pattern:
    """Generate ACARA code pattern for a subject.
    
    ACARA codes follow pattern: AC9{subject_code}{band}{strand_code}{sequence}
    Examples:
        Geography: AC9HG7K01 (subject=HG, band=7, strand=K, seq=01)
        Geography: AC9HG10S06 (subject=HG, band=10, strand=S, seq=06)
    """
    return re.compile(rf'AC9{subject_code}[F\d]{{1,2}}[A-Z]{{1,2}}\d{{2}}')


# Global pattern for pre-check validation (generic)
ACARA_CODE_PATTERN = re.compile(r'AC9[A-Z]{1,3}[F\d]{1,2}[A-Z]{1,2}\d{2}')


# =============================================================================
# STANDARDISED KEYWORD EXTRACTION (ALIGNED WITH VICTORIAN)
# =============================================================================

STOP_WORDS = {
    'the', 'and', 'of', 'to', 'a', 'an', 'in', 'on', 'at', 'for', 'with',
    'as', 'by', 'from', 'or', 'is', 'are', 'be', 'been', 'being',
    'that', 'this', 'these', 'those', 'they', 'them', 'their', 'it', 'its',
    'has', 'have', 'had', 'was', 'were', 'will', 'would', 'could', 'should',
    'may', 'might', 'must', 'shall', 'can', 'do', 'does', 'did',
    'how', 'which', 'also', 'than', 'then', 'end', 'both', 'each', 'such',
    'into', 'through', 'students', 'including', 'using'
}


# =============================================================================
# HISTORY TOPIC DETECTION
# =============================================================================

# Known topics for History subjects (ACARA and Victorian)
# These appear as headers within elaboration cells
HISTORY_TOPICS = {
    # Year 7 - The ancient world
    'Greece', 'Rome', 'Egypt', 'India', 'China',
    # Year 8 - Medieval Europe and the early modern world  
    'Medieval Europe', 'The Renaissance', 'The emergence of the modern world',
    # Year 8 - Empires and expansions
    'Mongol Empire', 'Ottoman Empire', 'Vikings', 'The Spanish conquest of the Americas',
    # Year 8 - Asia-Pacific world
    'Angkor/Khmer Empire', 'Japan under the Shoguns', 'Polynesian expansion across the Pacific',
}

# Verbs that start elaborations (NOT topics)
ELABORATION_VERBS = {
    'identifying', 'explaining', 'describing', 'examining', 'investigating',
    'analysing', 'analyzing', 'evaluating', 'comparing', 'creating',
    'developing', 'mapping', 'outlining', 'using', 'discussing',
    'reviewing', 'listening', 'brainstorming', 'ranking', 'collaborating',
    'locating', 'researching', 'exploring', 'considering', 'tracing'
}


def is_topic_line(line: str, subject: str) -> bool:
    """Detect if a line is a topic header rather than an elaboration.
    
    Topics are:
    - Short lines (< 60 characters)
    - Don't start with elaboration verbs
    - Start with uppercase
    - Match known topic patterns for History
    """
    if not line or len(line) > 60:
        return False
    
    # Must start with uppercase
    if line[0].islower():
        return False
    
    # Check if it starts with an elaboration verb
    first_word = line.split()[0].lower() if line.split() else ''
    if first_word in ELABORATION_VERBS:
        return False
    
    # For History subjects, check against known topics
    if subject == 'History':
        # Exact match
        if line in HISTORY_TOPICS:
            return True
        # Also check without "The " prefix
        if line.startswith('The ') and line[4:] in HISTORY_TOPICS:
            return True
    
    # Generic topic detection: short, capitalised, no elaboration verb
    # This catches topics not in the known list
    if len(line) < 40 and first_word not in ELABORATION_VERBS:
        # Additional heuristic: topics don't typically contain semicolons or "for example"
        if ';' not in line and 'for example' not in line.lower():
            return True
    
    return False


def extract_keywords(text: str) -> Set[str]:
    """Extract meaningful keywords. Min word length 3, no hard limit."""
    if not text:
        return set()
    words = text.lower().replace('.', '').replace(',', '').replace(';', '').split()
    return {w for w in words if len(w) > 3 and w not in STOP_WORDS}


# =============================================================================
# TEXT CLEANING (ALIGNED WITH VICTORIAN)
# =============================================================================

# US to AU English spelling conversions
US_TO_AU_SPELLING = {
    # -ize to -ise (and derivatives)
    'recognize': 'recognise', 'recognizes': 'recognises', 'recognized': 'recognised', 'recognizing': 'recognising',
    'organize': 'organise', 'organizes': 'organises', 'organized': 'organised', 'organizing': 'organising',
    'analyze': 'analyse', 'analyzes': 'analyses', 'analyzed': 'analysed', 'analyzing': 'analysing',
    'summarize': 'summarise', 'summarizes': 'summarises', 'summarized': 'summarised', 'summarizing': 'summarising',
    'categorize': 'categorise', 'categorizes': 'categorises', 'categorized': 'categorised', 'categorizing': 'categorising',
    'prioritize': 'prioritise', 'prioritizes': 'prioritises', 'prioritized': 'prioritised', 'prioritizing': 'prioritising',
    'utilize': 'utilise', 'utilizes': 'utilises', 'utilized': 'utilised', 'utilizing': 'utilising',
    'minimize': 'minimise', 'minimizes': 'minimises', 'minimized': 'minimised', 'minimizing': 'minimising',
    'maximize': 'maximise', 'maximizes': 'maximises', 'maximized': 'maximised', 'maximizing': 'maximising',
    'emphasize': 'emphasise', 'emphasizes': 'emphasises', 'emphasized': 'emphasised', 'emphasizing': 'emphasising',
    'visualize': 'visualise', 'visualizes': 'visualises', 'visualized': 'visualised', 'visualizing': 'visualising',
    'realize': 'realise', 'realizes': 'realises', 'realized': 'realised', 'realizing': 'realising',
    'normalize': 'normalise', 'normalizes': 'normalises', 'normalized': 'normalised', 'normalizing': 'normalising',
    'standardize': 'standardise', 'standardizes': 'standardises', 'standardized': 'standardised', 'standardizing': 'standardising',
    'customize': 'customise', 'customizes': 'customises', 'customized': 'customised', 'customizing': 'customising',
    'specialize': 'specialise', 'specializes': 'specialises', 'specialized': 'specialised', 'specializing': 'specialising',
    'generalize': 'generalise', 'generalizes': 'generalises', 'generalized': 'generalised', 'generalizing': 'generalising',
    'memorize': 'memorise', 'memorizes': 'memorises', 'memorized': 'memorised', 'memorizing': 'memorising',
    'theorize': 'theorise', 'theorizes': 'theorises', 'theorized': 'theorised', 'theorizing': 'theorising',
    'hypothesize': 'hypothesise', 'hypothesizes': 'hypothesises', 'hypothesized': 'hypothesised', 'hypothesizing': 'hypothesising',
    'synthesize': 'synthesise', 'synthesizes': 'synthesises', 'synthesized': 'synthesised', 'synthesizing': 'synthesising',
    'criticize': 'criticise', 'criticizes': 'criticises', 'criticized': 'criticised', 'criticizing': 'criticising',
    'apologize': 'apologise', 'apologizes': 'apologises', 'apologized': 'apologised', 'apologizing': 'apologising',
    # -or to -our
    'color': 'colour', 'colors': 'colours', 'colored': 'coloured', 'coloring': 'colouring',
    'behavior': 'behaviour', 'behaviors': 'behaviours',
    'favor': 'favour', 'favors': 'favours', 'favored': 'favoured', 'favoring': 'favouring', 'favorable': 'favourable',
    'honor': 'honour', 'honors': 'honours', 'honored': 'honoured', 'honoring': 'honouring',
    'labor': 'labour', 'labors': 'labours', 'labored': 'laboured', 'laboring': 'labouring',
    'neighbor': 'neighbour', 'neighbors': 'neighbours', 'neighboring': 'neighbouring',
    'humor': 'humour', 'humors': 'humours',
    'vigor': 'vigour',
    'endeavor': 'endeavour', 'endeavors': 'endeavours',
    # -er to -re
    'center': 'centre', 'centers': 'centres', 'centered': 'centred', 'centering': 'centring',
    'meter': 'metre', 'meters': 'metres',
    'theater': 'theatre', 'theaters': 'theatres',
    'fiber': 'fibre', 'fibers': 'fibres',
    # -og to -ogue
    'catalog': 'catalogue', 'catalogs': 'catalogues',
    'dialog': 'dialogue', 'dialogs': 'dialogues',
    'analog': 'analogue',
    'prolog': 'prologue',
    'epilog': 'epilogue',
    # Other common differences
    'program': 'programme', 'programs': 'programmes',
    'defense': 'defence',
    'offense': 'offence',
    'license': 'licence',
    'practice': 'practise',
    'judgment': 'judgement',
    'aging': 'ageing',
    'modeling': 'modelling',
    'traveling': 'travelling', 'traveled': 'travelled', 'traveler': 'traveller',
    'canceled': 'cancelled', 'canceling': 'cancelling',
    'labeled': 'labelled', 'labeling': 'labelling',
    'leveled': 'levelled', 'leveling': 'levelling',
    'signaled': 'signalled', 'signaling': 'signalling',
}


def convert_to_au_english(text: str) -> str:
    """Convert US English spellings to Australian English."""
    if not text:
        return text
    
    result = text
    for us_spelling, au_spelling in US_TO_AU_SPELLING.items():
        pattern = re.compile(re.escape(us_spelling), re.IGNORECASE)
        
        def replace_preserve_case(match):
            matched = match.group(0)
            if matched.isupper():
                return au_spelling.upper()
            elif matched[0].isupper():
                return au_spelling.capitalize()
            return au_spelling
        
        result = pattern.sub(replace_preserve_case, result)
    
    return result


def capitalise_first(text: str) -> str:
    """Capitalise first letter, preserve rest."""
    if not text:
        return ""
    return text[0].upper() + text[1:] if len(text) > 1 else text.upper()


def clean_text(text: str) -> str:
    """Clean text while preserving newlines and converting to AU English."""
    if not text:
        return ""
    text = re.sub(r'[^\S\n]+', ' ', text)
    lines = text.split('\n')
    cleaned = '\n'.join(line.strip() for line in lines if line.strip())
    return convert_to_au_english(cleaned)


def normalise_acara_band(band: str) -> str:
    """Normalise ACARA band names to match VCAA format.
    
    ACARA documents use "Years X–Y" (with en-dash).
    VCAA uses "Years X and Y" format.
    This ensures consistency across both curricula.
    
    Examples:
        "Years 1–2"  -> "Years 1 and 2"
        "Years 7–8"  -> "Years 7 and 8"
        "Years 9–10" -> "Years 9 and 10"
        "Foundation" -> "Foundation"
        "Year 7"     -> "Year 7"
    """
    if not band:
        return band
    
    band = band.strip()
    
    # Replace en-dash, em-dash, or hyphen between numbers with " and "
    # Pattern: Years followed by number, dash, number
    normalised = re.sub(r'^(Years?\s*)(\d+)\s*[–—-]\s*(\d+)$', r'Years \2 and \3', band)
    
    return normalised


def map_as_year_to_band(year_text: str, subject_area: str) -> str:
    """Map Achievement Standard endpoint year to actual band name.
    
    For Technologies and The Arts subjects, AS text says "By the end of Year 2" but
    the actual band is "Years 1 and 2".
    
    Mapping for Technologies (F-10):
        "Foundation" -> "Foundation"
        "Year 2"     -> "Years 1 and 2"
        "Year 4"     -> "Years 3 and 4"
        "Year 6"     -> "Years 5 and 6"
        "Year 8"     -> "Years 7 and 8"
        "Year 10"    -> "Years 9 and 10"
    
    Mapping for The Arts (7-10 only - F-6 uses single year bands):
        "Year 8"     -> "Years 7 and 8"
        "Year 10"    -> "Years 9 and 10"
    
    For other subjects (HASS, English, etc.), returns unchanged.
    """
    if subject_area == 'Technologies':
        if year_text == 'Foundation':
            return 'Foundation'
        
        tech_band_map = {
            'Year 2': 'Years 1 and 2',
            'Year 4': 'Years 3 and 4',
            'Year 6': 'Years 5 and 6',
            'Year 8': 'Years 7 and 8',
            'Year 10': 'Years 9 and 10',
        }
        return tech_band_map.get(year_text, year_text)
    
    elif subject_area == 'The Arts':
        # The Arts uses banded structure for all year levels
        # F-6: Years 1 and 2, Years 3 and 4, Years 5 and 6
        # 7-10: Years 7 and 8, Years 9 and 10
        arts_band_map = {
            'Year 2': 'Years 1 and 2',
            'Year 4': 'Years 3 and 4',
            'Year 6': 'Years 5 and 6',
            'Year 8': 'Years 7 and 8',
            'Year 10': 'Years 9 and 10',
        }
        return arts_band_map.get(year_text, year_text)
    
    return year_text


def normalise_band_name(band: str) -> str:
    """Normalise band names to consistent format.
    
    VCAA uses "Years X and Y" format for multi-year bands.
    This function ensures ACARA bands follow the same convention.
    
    Conversions:
    - "Years 7-8" -> "Years 7 and 8"
    - "Years 9 - 10" -> "Years 9 and 10"  
    - "Year 7-8" -> "Years 7 and 8"
    - "Levels 3-4" -> "Levels 3 and 4"
    - "Years 7 and 8" -> unchanged
    - "Year 7" -> unchanged
    - "Foundation" -> unchanged
    """
    if not band:
        return band
    
    # Pattern for "Years/Year/Levels X-Y" or "Years/Year/Levels X - Y"
    pattern = r'^(Years?|Levels?)\s*(\d+)\s*[-–]\s*(\d+)$'
    match = re.match(pattern, band, re.IGNORECASE)
    if match:
        prefix = match.group(1)
        start = match.group(2)
        end = match.group(3)
        # Ensure prefix is plural for ranges
        if prefix.lower() in ['year', 'level']:
            prefix = prefix + 's'
        return f"{prefix} {start} and {end}"
    
    return band


def get_cell_text(cell) -> str:
    """Extract text from cell, handling nested tables."""
    cell_text = cell.text.strip()
    if not cell_text:
        from docx.table import Table
        nested = cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl')
        if nested:
            texts = []
            for tbl_el in nested:
                nt = Table(tbl_el, cell._parent._parent._parent)
                for row in nt.rows:
                    for nc in row.cells:
                        t = nc.text.strip()
                        if t:
                            texts.append(t)
            cell_text = '\n'.join(texts)
    return cell_text


# =============================================================================
# DATA CLASSES
# =============================================================================

@dataclass
class ContentDescriptor:
    code: str
    description: str
    subject_area: str
    subject: str
    band: str
    strand: str
    substrand: str
    topic: str = ""  # Default topic (for non-History subjects)
    elaborations: List[Tuple[str, str]] = field(default_factory=list)  # List of (topic, elaboration) pairs
    eald_elaborations: List[str] = field(default_factory=list)

@dataclass
class AchievementStandard:
    subject_area: str
    subject: str
    band: str
    text: str

@dataclass
class ASComponent:
    code: str
    subject_area: str
    subject: str
    band: str
    text: str
    strand: str
    keywords: str
    linked_codes: str
    confidence: str

@dataclass
class GlossaryTerm:
    subject_area: str
    subject: str
    term: str
    definition: str


# =============================================================================
# PRE-CHECK TESTS
# =============================================================================

class ACARAPrecheckTests:
    """Pre-parsing validation tests."""
    
    def __init__(self, curriculum_docs, eald_docs=None, glossary_doc=None):
        self.curriculum_docs = curriculum_docs
        self.eald_docs = eald_docs or []
        self.glossary_doc = glossary_doc
        self.issues = []
        self.warnings = []
        self.passed_tests = 0
        self.total_tests = 0
    
    def run_all(self) -> bool:
        print("\n" + "="*80)
        print("PHASE 1: PRE-CHECK VALIDATION")
        print("="*80)
        
        self._check_document_access()
        self._check_document_structure()
        self._check_for_nested_tables()
        self._check_code_patterns()
        
        print("\n" + "-"*80)
        print(f"Tests passed: {self.passed_tests}/{self.total_tests}")
        if self.issues:
            print(f"\n❌ CRITICAL: {self.issues}")
        if self.warnings:
            print(f"\n⚠️  WARNINGS: {self.warnings}")
        print("="*80)
        
        return len(self.issues) == 0
    
    def _check_document_access(self):
        self.total_tests += 1
        print("\n1. DOCUMENT ACCESS")
        all_docs = self.curriculum_docs + self.eald_docs
        if self.glossary_doc:
            all_docs.append(self.glossary_doc)
        
        accessible = 0
        for doc_path in all_docs:
            try:
                Document(doc_path)
                print(f"   ✅ {doc_path.split('/')[-1]}")
                accessible += 1
            except Exception as e:
                self.issues.append(f"Cannot access {doc_path}")
                print(f"   ❌ {doc_path.split('/')[-1]}: ERROR")
        
        if accessible == len(all_docs):
            self.passed_tests += 1
    
    def _check_document_structure(self):
        self.total_tests += 1
        print("\n2. DOCUMENT STRUCTURE")
        all_valid = True
        for doc_path in self.curriculum_docs:
            try:
                doc = Document(doc_path)
                print(f"   {doc_path.split('/')[-1]}: {len(doc.tables)} tables, {len(doc.paragraphs)} paras")
                if len(doc.tables) == 0:
                    self.issues.append(f"No tables in {doc_path}")
                    all_valid = False
            except Exception:
                all_valid = False
        if all_valid:
            self.passed_tests += 1
    
    def _check_for_nested_tables(self):
        self.total_tests += 1
        print("\n3. NESTED TABLES SCAN")
        nested_count = 0
        for doc_path in self.curriculum_docs:
            try:
                doc = Document(doc_path)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            nested = cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl')
                            nested_count += len(nested)
            except Exception:
                pass
        
        if nested_count > 0:
            print(f"   ℹ️  {nested_count} nested tables (get_cell_text will handle)")
        else:
            print("   ✅ No nested tables")
        self.passed_tests += 1
    
    def _check_code_patterns(self):
        self.total_tests += 1
        print("\n4. ACARA CODE PATTERN CHECK")
        codes_found = set()
        for doc_path in self.curriculum_docs:
            try:
                doc = Document(doc_path)
                for table in doc.tables[:5]:
                    for row in table.rows:
                        for cell in row.cells:
                            codes_found.update(ACARA_CODE_PATTERN.findall(cell.text))
            except Exception:
                pass
        
        if codes_found:
            print(f"   ✅ Found {len(codes_found)} unique ACARA codes")
            self.passed_tests += 1
        else:
            self.warnings.append("No ACARA codes found")
            print("   ⚠️  No ACARA codes detected")


# =============================================================================
# POST-PARSE TESTS
# =============================================================================

class ACARAParsedDataTests:
    """Tests run after parsing."""
    
    def __init__(self, curriculum_data, achievement_standards, as_components, glossary_terms, code_pattern):
        self.curriculum_data = curriculum_data
        self.achievement_standards = achievement_standards
        self.as_components = as_components
        self.glossary_terms = glossary_terms
        self.code_pattern = code_pattern
        self.test_results = []
    
    def run_all(self) -> bool:
        print("\n" + "="*80)
        print("PHASE 3: POST-PARSE DATA VALIDATION")
        print("="*80)
        
        self._test_code_extraction()
        self._test_capitalisation()
        self._test_strand_inference()
        self._test_elaboration_counts()
        self._test_topic_extraction()
        self._test_no_duplicates()
        
        print("\n" + "-"*80)
        passed = sum(1 for r in self.test_results if r['passed'])
        print(f"Total: {passed}/{len(self.test_results)} tests passed")
        print("="*80)
        
        return all(r['passed'] for r in self.test_results)
    
    def _test_code_extraction(self):
        print("\n1. CODE EXTRACTION")
        invalid = [d['code'] for d in self.curriculum_data if not self.code_pattern.match(d.get('code', ''))]
        passed = len(invalid) == 0
        self.test_results.append({'test': 'Code Extraction', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} {len(self.curriculum_data)} codes, {len(invalid)} invalid")
        if invalid:
            print(f"      Invalid: {invalid[:5]}")
    
    def _test_capitalisation(self):
        print("\n2. CAPITALISATION")
        uncap_desc = [d for d in self.curriculum_data if d.get('description', '') and d['description'][0].islower()]
        # Handle elaborations as tuples (topic, text) - check the text part
        uncap_elab = 0
        for d in self.curriculum_data:
            for elab in d.get('elaborations', []):
                # elab is either a tuple (topic, text) or a string
                text = elab[1] if isinstance(elab, tuple) else elab
                if text and text[0].islower():
                    uncap_elab += 1
        passed = len(uncap_desc) == 0 and uncap_elab == 0
        self.test_results.append({'test': 'Capitalisation', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} Descriptions: {len(uncap_desc)} uncap, Elaborations: {uncap_elab}")
    
    def _test_strand_inference(self):
        print("\n3. STRAND INFERENCE")
        no_strand = [c for c in self.as_components if not c.get('strand')]
        passed = len(no_strand) == 0
        self.test_results.append({'test': 'Strand Inference', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} {len(self.as_components)} components, {len(no_strand)} missing strand")
    
    def _test_elaboration_counts(self):
        print("\n4. ELABORATIONS")
        total = sum(len(d.get('elaborations', [])) for d in self.curriculum_data)
        with_elab = sum(1 for d in self.curriculum_data if d.get('elaborations'))
        passed = total > 0
        self.test_results.append({'test': 'Elaborations', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} {total} total, {with_elab} CDs have elaborations")
    
    def _test_topic_extraction(self):
        print("\n5. TOPIC EXTRACTION")
        # Count elaborations with topics
        with_topics = 0
        unique_topics = set()
        for d in self.curriculum_data:
            for elab in d.get('elaborations', []):
                if isinstance(elab, tuple) and elab[0]:  # (topic, text) tuple with non-empty topic
                    with_topics += 1
                    unique_topics.add(elab[0])
        
        # This test always passes - it's informational
        passed = True
        self.test_results.append({'test': 'Topic Extraction', 'passed': passed})
        if with_topics > 0:
            print(f"   ℹ️  {with_topics} elaborations with topics, {len(unique_topics)} unique topics")
            if unique_topics:
                print(f"      Topics: {', '.join(sorted(unique_topics)[:10])}" + 
                      ("..." if len(unique_topics) > 10 else ""))
        else:
            print(f"   ℹ️  No embedded topics detected (normal for non-History subjects)")
    
    def _test_no_duplicates(self):
        print("\n6. DUPLICATES")
        codes = [d.get('code') for d in self.curriculum_data]
        dups = [c for c in set(codes) if codes.count(c) > 1]
        passed = len(dups) == 0
        self.test_results.append({'test': 'Duplicates', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} {len(set(codes))} unique, {len(dups)} duplicates")


# =============================================================================
# PARSER CLASS
# =============================================================================

class ACARACurriculumParser:
    """Generic parser for ACARA curriculum documents."""
    
    def __init__(self, config: dict):
        self.config = config
        self.content_descriptors: List[ContentDescriptor] = []
        self.achievement_standards: List[AchievementStandard] = []
        self.as_components: List[ASComponent] = []
        self.glossary_terms: List[GlossaryTerm] = []
        self.discovered_strands: Dict[str, List[str]] = defaultdict(list)
        
        # Generate subject-specific code pattern
        self.subject_code = config.get('subject_code', config['subject'][0].upper())
        self.code_pattern = get_acara_code_pattern(self.subject_code)
        self.strand_codes = config.get('strand_codes', {})
        
        os.makedirs(config['output_dir'], exist_ok=True)
        
        print(f"ACARA Curriculum Parser v{PARSER_VERSION}")
        print(f"Subject: {config['subject']} (code: {self.subject_code})")
        print(f"Output: {config['output_dir']}")
        print("=" * 60)
    
    def extract_acara_code(self, text: str) -> Optional[str]:
        match = self.code_pattern.search(text)
        return match.group(0) if match else None
    
    def get_strand_from_code(self, code: str) -> str:
        # Extract strand code from pattern AC9{subject}{band}{strand_code}{seq}
        pattern = rf'AC9{self.subject_code}[F\d]{{1,2}}([A-Z]{{1,2}})\d{{2}}'
        match = re.search(pattern, code)
        if match:
            strand_code = match.group(1)
            return self.strand_codes.get(strand_code, strand_code)
        return ""
    
    def get_band_from_code(self, code: str) -> str:
        pattern = rf'AC9{self.subject_code}([F\d]{{1,2}})[A-Z]{{1,2}}\d{{2}}'
        match = re.search(pattern, code)
        if match:
            year_part = match.group(1)
            return 'Foundation' if year_part == 'F' else f'Year {year_part}'
        return ""
    
    def clean_description(self, text: str) -> str:
        text = self.code_pattern.sub('', text)
        return clean_text(text).strip()
    
    def parse_curriculum_document(self, doc_path: str, bands: List[str]) -> None:
        print(f"\nParsing: {os.path.basename(doc_path)}")
        doc = Document(doc_path)
        
        # For Technologies subjects, we need to track band changes from description tables
        is_technologies = self.config.get('subject_area') == 'Technologies'
        current_band = None
        
        for table in doc.tables:
            if len(table.rows) < 1:
                continue
            
            first_row_text = get_cell_text(table.rows[0].cells[0])
            
            # Detect band changes from description tables (Technologies subjects)
            if is_technologies and ('Year level description' in first_row_text or 'Band level description' in first_row_text):
                if len(table.rows) >= 4:
                    # Band is in the AS text which starts with "By the end of [Band]"
                    as_row_text = get_cell_text(table.rows[3].cells[0])
                    band_match = re.search(r'By the end of (Foundation|Year \d+)', as_row_text)
                    if band_match:
                        current_band = map_as_year_to_band(band_match.group(1), 'Technologies')
                        print(f"   Band change detected: {current_band}")
                continue  # Skip to next table
            
            # Skip other non-curriculum tables
            if 'Achievement standard' in first_row_text:
                continue
            if 'Learning area' in first_row_text:
                continue
            
            # Process content table with current band context
            self._process_content_table(table, bands, current_band if is_technologies else None)
    
    def _process_content_table(self, table, bands: List[str], inherited_band: str = None) -> None:
        current_strand = None
        current_substrand = None
        current_band = inherited_band  # Use inherited band from Technologies description tables
        
        for row in table.rows:
            cells = row.cells
            if not cells:
                continue
            
            first_cell_text = get_cell_text(cells[0])
            
            if first_cell_text.startswith('Strand:'):
                current_strand = first_cell_text.replace('Strand:', '').strip()
                if len(cells) > 1:
                    last_cell = normalise_acara_band(cells[-1].text.strip())
                    if last_cell in bands:
                        current_band = last_cell
                continue
            
            if first_cell_text.startswith('Sub-strand:'):
                current_substrand = first_cell_text.replace('Sub-strand:', '').strip()
                continue
            
            if 'Content descriptions' in first_cell_text or 'Students learn to' in first_cell_text or 'Students learn about' in first_cell_text:
                continue
            
            code = self.extract_acara_code(first_cell_text)
            if code:
                description = capitalise_first(self.clean_description(first_cell_text))
                
                if not current_band:
                    current_band = self.get_band_from_code(code)
                if not current_strand:
                    current_strand = self.get_strand_from_code(code)
                
                # Check for duplicate before adding - use raw elaboration text for comparison
                existing = next((cd for cd in self.content_descriptors if cd.code == code), None)
                if existing:
                    # Already processed this code, skip
                    continue
                
                # Parse elaborations with topic detection
                elaborations = []  # List of (topic, elaboration) tuples
                seen_elaborations = set()  # Track raw text before capitalisation
                current_topic = ""  # Track current topic for History subjects
                
                for cell_idx in range(1, len(cells)):
                    cell_text = get_cell_text(cells[cell_idx])
                    if cell_text and 'Content elaborations' not in cell_text:
                        for line in cell_text.split('\n'):
                            raw_line = line.strip()
                            if not raw_line:
                                continue
                            
                            # Check if this line is a topic header
                            if is_topic_line(raw_line, self.config['subject']):
                                current_topic = raw_line
                                continue  # Don't add topic as an elaboration
                            
                            # It's an elaboration - pair with current topic
                            if raw_line not in seen_elaborations:
                                seen_elaborations.add(raw_line)
                                elaborations.append((current_topic, capitalise_first(raw_line)))
                
                cd = ContentDescriptor(
                    code=code,
                    description=description,
                    subject_area=self.config['subject_area'],
                    subject=self.config['subject'],
                    band=current_band,
                    strand=current_strand or self.get_strand_from_code(code),
                    substrand=current_substrand or "",
                    elaborations=elaborations
                )
                self.content_descriptors.append(cd)
                
                if current_strand and current_substrand:
                    if current_substrand not in self.discovered_strands[current_strand]:
                        self.discovered_strands[current_strand].append(current_substrand)
    
    def parse_achievement_standards(self, doc_path: str, bands: List[str]) -> None:
        """Parse achievement standards from curriculum document.
        
        TECHNOLOGIES SUBJECTS (subject_area == 'Technologies'):
        These subjects have TWO achievement standards in documents:
        - Subject-specific AS (Row 3) - shorter, less detailed
        - Learning area Achievement standard (Row 5) - comprehensive, PREFERRED
        
        F-6 Table structure for Technologies:
        - Row 0: 'Year level description' or 'Band level description'
        - Row 1: Description text
        - Row 2: '[Subject] Achievement standard' (header)
        - Row 3: Short subject-specific AS (SKIP for Technologies)
        - Row 4: 'Learning area Achievement standard' (header)
        - Row 5: Full AS text (USE THIS for Technologies)
        
        7-10 Table structure for Technologies (varies):
        - Some bands: Standalone 'Learning area Achievement standard' table
        - Some bands: Only has 'Achievement standard' in Band level description table
        
        OTHER SUBJECTS (HASS, English, etc.):
        Use the standard Row 3 AS from 'Year level description' tables.
        """
        print(f"Extracting achievement standards from: {os.path.basename(doc_path)}")
        doc = Document(doc_path)
        
        # Check if this is a Technologies subject
        is_technologies = self.config.get('subject_area') == 'Technologies'
        
        if is_technologies:
            self._parse_technologies_achievement_standards(doc)
        else:
            self._parse_standard_achievement_standards(doc)
    
    def _parse_technologies_achievement_standards(self, doc: Document) -> None:
        """Parse AS for Technologies subjects using Learning area Achievement standard.
        
        Technologies subjects use the 'Learning area Achievement standard' which is
        more comprehensive than the subject-specific AS.
        """
        # PASS 1: Process Learning area AS (preferred)
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            
            first_row_text = get_cell_text(table.rows[0].cells[0])
            
            # Option A: Learning area AS embedded in description tables (Row 5)
            if ('Year level description' in first_row_text or 'Band level description' in first_row_text) and len(table.rows) >= 6:
                row4_text = get_cell_text(table.rows[4].cells[0]).lower()
                if 'learning area' in row4_text and 'achievement standard' in row4_text:
                    as_text = clean_text(get_cell_text(table.rows[5].cells[0]))
                    band_match = re.search(r'By the end of (Foundation|Year \d+)', as_text)
                    if band_match:
                        band = map_as_year_to_band(band_match.group(1), 'Technologies')
                        if not any(a.band == band for a in self.achievement_standards):
                            self.achievement_standards.append(AchievementStandard(
                                self.config['subject_area'], self.config['subject'], band, as_text))
                            print(f"   Found AS for: {band} (Learning area)")
            
            # Option B: Standalone 'Learning area Achievement standard' table
            elif 'Learning area Achievement standard' in first_row_text:
                as_text = clean_text(get_cell_text(table.rows[1].cells[0]))
                band_match = re.search(r'By the end of (Foundation|Year \d+)', as_text)
                if band_match:
                    band = map_as_year_to_band(band_match.group(1), 'Technologies')
                    if not any(a.band == band for a in self.achievement_standards):
                        self.achievement_standards.append(AchievementStandard(
                            self.config['subject_area'], self.config['subject'], band, as_text))
                        print(f"   Found AS for: {band} (Learning area standalone)")
        
        # PASS 2: Fallback for bands without Learning area AS
        for table in doc.tables:
            if len(table.rows) < 4:
                continue
            
            first_row_text = get_cell_text(table.rows[0].cells[0])
            
            if 'Band level description' in first_row_text:
                # Only use fallback if this table doesn't have Learning area AS
                has_learning_area = len(table.rows) >= 6 and 'learning area' in get_cell_text(table.rows[4].cells[0]).lower()
                if not has_learning_area:
                    row2_text = get_cell_text(table.rows[2].cells[0]).lower()
                    if 'achievement standard' in row2_text:
                        as_text = clean_text(get_cell_text(table.rows[3].cells[0]))
                        band_match = re.search(r'By the end of (Foundation|Year \d+)', as_text)
                        if band_match:
                            band = map_as_year_to_band(band_match.group(1), 'Technologies')
                            # Only add if not already captured via Learning area
                            if not any(a.band == band for a in self.achievement_standards):
                                self.achievement_standards.append(AchievementStandard(
                                    self.config['subject_area'], self.config['subject'], band, as_text))
                                print(f"   Found AS for: {band} (fallback)")
    
    def _parse_standard_achievement_standards(self, doc: Document) -> None:
        """Parse AS for standard subjects (HASS, English, The Arts, etc.) using Row 3.
        
        Handles both 'Year level description' (most subjects) and 'Band level description'
        (The Arts 7-10 documents) table structures.
        """
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            
            first_row_text = get_cell_text(table.rows[0].cells[0])
            
            # Check for BOTH 'Year level description' AND 'Band level description'
            # The Arts 7-10 documents use 'Band level description'
            if ('Year level description' in first_row_text or 'Band level description' in first_row_text) and len(table.rows) >= 4:
                as_text = clean_text(get_cell_text(table.rows[3].cells[0]))
                # Match both "By the end of Foundation" and "By the end of the Foundation year"
                band_match = re.search(r'By the end of (?:the )?(Foundation(?: year)?|Year \d+)', as_text)
                if band_match:
                    # Extract just "Foundation" or "Year X" from the match
                    matched_text = band_match.group(1).replace(' year', '')
                    # Apply band normalisation for consistency (Year 8 -> Years 7 and 8)
                    band = map_as_year_to_band(matched_text, self.config.get('subject_area', ''))
                    if not any(a.band == band for a in self.achievement_standards):
                        self.achievement_standards.append(AchievementStandard(
                            self.config['subject_area'], self.config['subject'], band, as_text))
            
            elif 'Achievement standard' in first_row_text:
                as_text = clean_text(get_cell_text(table.rows[1].cells[0]))
                # Match both "By the end of Foundation" and "By the end of the Foundation year"
                band_match = re.search(r'By the end of (?:the )?(Foundation(?: year)?|Year \d+)', as_text)
                if band_match:
                    # Extract just "Foundation" or "Year X" from the match
                    matched_text = band_match.group(1).replace(' year', '')
                    band = map_as_year_to_band(matched_text, self.config.get('subject_area', ''))
                    if not any(a.band == band for a in self.achievement_standards):
                        self.achievement_standards.append(AchievementStandard(
                            self.config['subject_area'], self.config['subject'], band, as_text))
    
    def generate_as_components(self) -> None:
        """Generate AS components using generic keyword-based strand inference."""
        print("\nGenerating AS Components...")
        
        # Track component count per band for code generation
        band_component_counts = defaultdict(int)
        
        for ast in self.achievement_standards:
            sentences = self._split_sentences(ast.text)
            for sentence in sentences:
                if len(sentence) < 20:
                    continue
                
                # Generic strand inference using keyword matching
                strand = self._infer_strand_generic(sentence, ast.band)
                keywords = extract_keywords(sentence)
                linked_codes, confidence = self._find_linked_codes(ast.band, strand, keywords)
                
                # Generate ASComponentCode: AC9{subject_code}{band_code}ASC{sequence:02d}
                band_code = self._get_band_code(ast.band)
                band_component_counts[ast.band] += 1
                asc_code = f"AC9{self.subject_code}{band_code}ASC{band_component_counts[ast.band]:02d}"
                
                self.as_components.append(ASComponent(
                    code=asc_code,
                    subject_area=ast.subject_area,
                    subject=ast.subject,
                    band=ast.band,
                    text=sentence,
                    strand=strand,
                    keywords=', '.join(sorted(keywords)),
                    linked_codes=', '.join(linked_codes),
                    confidence=confidence
                ))
    
    def _get_band_code(self, band: str) -> str:
        """Convert band name to code for ASComponentCode generation.
        
        Examples:
            'Foundation'      -> 'F'
            'Year 7'          -> '7'
            'Years 1 and 2'   -> '2'   (use end year)
            'Years 7 and 8'   -> '8'   (use end year)
        """
        if band == 'Foundation':
            return 'F'
        
        # Handle "Years X and Y" format - use the end year
        match = re.match(r'Years? (\d+) and (\d+)', band)
        if match:
            return match.group(2)  # Return end year
        
        # Handle "Year X" format
        match = re.match(r'Year (\d+)', band)
        if match:
            return match.group(1)
        
        return band[0].upper()
    
    def _split_sentences(self, text: str) -> List[str]:
        paragraphs = text.split('\n')
        sentences = []
        for para in paragraphs:
            parts = re.split(r'(?<=[.!?])\s+(?=[A-Z])', para)
            for part in parts:
                part = part.strip()
                if part:
                    sentences.append(part)
        return sentences
    
    def _infer_strand_generic(self, text: str, band: str) -> str:
        """Generic strand inference using keyword overlap with content descriptors."""
        if not self.content_descriptors:
            return "General"
        
        component_keywords = extract_keywords(text.lower())
        strand_scores = {}
        
        for cd in self.content_descriptors:
            if cd.band != band:
                continue
            cd_keywords = extract_keywords(cd.description.lower())
            overlap = len(component_keywords & cd_keywords)
            if cd.strand not in strand_scores:
                strand_scores[cd.strand] = 0
            strand_scores[cd.strand] += overlap
        
        if strand_scores and max(strand_scores.values()) > 0:
            return max(strand_scores, key=strand_scores.get)
        return "General"
    
    def _find_linked_codes(self, band: str, strand: str, keywords: Set[str]) -> Tuple[List[str], str]:
        matched = []
        for cd in self.content_descriptors:
            if cd.band != band:
                continue
            score = 2 if cd.strand == strand else 0
            # Extract elaboration text from tuples
            elab_texts = [elab for _, elab in cd.elaborations] if cd.elaborations else []
            cd_text = (cd.description + ' ' + ' '.join(elab_texts)).lower()
            score += sum(1 for kw in keywords if kw in cd_text)
            if score >= 2:
                matched.append((cd.code, score))
        
        matched.sort(key=lambda x: x[1], reverse=True)
        top_codes = [code for code, _ in matched[:5]]
        
        if matched and matched[0][1] >= 4:
            confidence = "High"
        elif matched and matched[0][1] >= 2:
            confidence = "Medium"
        else:
            confidence = "Low"
        
        return top_codes, confidence
    
    def parse_glossary(self, doc_path: str) -> None:
        print(f"\nParsing glossary: {os.path.basename(doc_path)}")
        doc = Document(doc_path)
        
        # Track seen terms to avoid duplicates
        seen_terms = set()
        
        for table in doc.tables:
            for row in table.rows:
                text = get_cell_text(row.cells[0])
                if len(text) <= 2 or 'Glossary' in text:
                    continue
                # Skip single letter rows (alphabet headers like "A", "B", etc.)
                if len(text.strip()) == 1:
                    continue
                
                term = None
                definition = None
                
                # Try tab separator first (most common)
                if '\t' in text:
                    parts = text.split('\t', 1)
                    term = parts[0].strip()
                    # Handle case where tab is followed by newline (e.g., "Country/Place\t\n...")
                    if len(parts) > 1:
                        def_text = parts[1].strip()
                        if def_text:
                            definition = clean_text(def_text)
                        elif '\n' in text:
                            # Tab was empty, try newline for definition
                            newline_parts = text.split('\n', 1)
                            if len(newline_parts) > 1:
                                definition = clean_text(newline_parts[1])
                
                # If no tab, try newline separator (for terms with slashes like "Asia / Asian")
                elif '\n' in text:
                    parts = text.split('\n', 1)
                    term = parts[0].strip()
                    definition = clean_text(parts[1]) if len(parts) > 1 else ""
                
                if term and definition and term.lower() not in seen_terms:
                    seen_terms.add(term.lower())
                    # Capitalise first letter of term for consistency
                    term = capitalise_first(term)
                    self.glossary_terms.append(GlossaryTerm(
                        self.config['subject_area'], self.config['subject'], term, definition))
        
        print(f"Extracted {len(self.glossary_terms)} terms")
    
    # -------------------------------------------------------------------------
    # CSV OUTPUT
    # -------------------------------------------------------------------------
    
    def write_curriculum_csvs(self) -> None:
        print("\nWriting curriculum CSVs...")
        by_band = defaultdict(list)
        for cd in self.content_descriptors:
            by_band[cd.band].append(cd)
        
        for band, cds in by_band.items():
            filename = f"ACARA {self.config['subject']} V9 - Curriculum - {band}.csv"
            filepath = os.path.join(self.config['output_dir'], filename)
            
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow([
                    'SubjectArea', 'Subject', 'Band', 'Strand', 'Substrand',
                    'ContentCode', 'ContentDescription', 'Topic', 'Elaboration', 'EALDElaboration'
                ])
                
                for cd in cds:
                    num_regular = len(cd.elaborations) if cd.elaborations else 1
                    num_eald = len(cd.eald_elaborations)
                    max_rows = max(num_regular, 1)
                    
                    for i in range(max_rows):
                        # Get topic and elaboration from tuple (topic, elab_text)
                        topic = ""
                        elab = ""
                        if cd.elaborations and i < len(cd.elaborations):
                            topic, elab = cd.elaborations[i]
                        
                        # Fall back to cd.topic if no topic from tuple (for non-History subjects)
                        if not topic and cd.topic:
                            topic = cd.topic
                        
                        eald = ""
                        if i < num_eald:
                            if i == 0:
                                eald = f"EAL/D students may benefit from:\n\n{cd.eald_elaborations[i]}"
                            else:
                                eald = cd.eald_elaborations[i]
                        
                        writer.writerow([
                            cd.subject_area, cd.subject, cd.band, cd.strand, cd.substrand,
                            cd.code, cd.description, topic, elab, eald
                        ])
            
            print(f"  {filename}: {len(cds)} descriptors")
    
    def write_achievement_standards_csv(self) -> None:
        filename = f"ACARA {self.config['subject']} V9 - Achievement Standards.csv"
        filepath = os.path.join(self.config['output_dir'], filename)
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['SubjectArea', 'Subject', 'Band', 'AchievementStandard'])
            
            all_bands = self.config['bands_f6'] + self.config['bands_7_10']
            sorted_as = sorted(self.achievement_standards,
                             key=lambda x: all_bands.index(x.band) if x.band in all_bands else 99)
            
            for ast in sorted_as:
                writer.writerow([ast.subject_area, ast.subject, ast.band, ast.text])
        
        print(f"\n{filename}: {len(self.achievement_standards)} standards")
    
    def write_as_components_csv(self) -> None:
        filename = f"ACARA {self.config['subject']} V9 - AS Components.csv"
        filepath = os.path.join(self.config['output_dir'], filename)
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            # Column order matches import patterns skill
            writer.writerow([
                'SubjectArea', 'Subject', 'Band', 'ASComponentCode', 'ASComponentText',
                'Strand', 'Keywords', 'LinkedContentCodes', 'ConfidenceScore'
            ])
            for comp in self.as_components:
                writer.writerow([
                    comp.subject_area, comp.subject, comp.band, comp.code, comp.text,
                    comp.strand, comp.keywords, comp.linked_codes, comp.confidence
                ])
        
        print(f"{filename}: {len(self.as_components)} components")
    
    def write_glossary_csv(self) -> None:
        filename = f"ACARA {self.config['subject']} V9 - Glossary.csv"
        filepath = os.path.join(self.config['output_dir'], filename)
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            # Include Band column to match import patterns skill
            writer.writerow(['SubjectArea', 'Subject', 'Band', 'Term', 'Definition'])
            for gt in sorted(self.glossary_terms, key=lambda x: x.term.lower()):
                # F-10 glossaries use "Foundation to Year 10" band
                writer.writerow([gt.subject_area, gt.subject, 'Foundation to Year 10', gt.term, gt.definition])
        
        print(f"{filename}: {len(self.glossary_terms)} terms")
    
    # -------------------------------------------------------------------------
    # VALIDATION
    # -------------------------------------------------------------------------
    
    def validate(self) -> bool:
        print("\n" + "=" * 60)
        print("VALIDATION")
        print("=" * 60)
        
        print(f"Content descriptors: {len(self.content_descriptors)}")
        print(f"Achievement standards: {len(self.achievement_standards)}")
        print(f"AS components: {len(self.as_components)}")
        print(f"Glossary terms: {len(self.glossary_terms)}")
        
        codes = [cd.code for cd in self.content_descriptors]
        dups = [c for c in set(codes) if codes.count(c) > 1]
        if dups:
            print(f"⚠️  Duplicate codes: {dups}")
        
        invalid = [cd.code for cd in self.content_descriptors if not self.code_pattern.match(cd.code)]
        if invalid:
            print(f"⚠️  Invalid codes: {invalid}")
        
        return len(dups) == 0 and len(invalid) == 0
    
    # -------------------------------------------------------------------------
    # MAIN EXECUTION
    # -------------------------------------------------------------------------
    
    def run(self) -> None:
        """Execute full parsing pipeline."""
        
        # Phase 1: Pre-check
        precheck = ACARAPrecheckTests(
            curriculum_docs=self.config['curriculum_docs'],
            eald_docs=self.config.get('eald_docs', []),
            glossary_doc=self.config.get('glossary_doc')
        )
        
        if not precheck.run_all():
            print("\n❌ CANCELLED - Fix issues first")
            return
        
        print("\n✅ Pre-checks passed\n")
        
        # Phase 2: Parse
        print("=" * 60)
        print("PARSING CURRICULUM DOCUMENTS")
        print("=" * 60)
        
        # Parse all curriculum documents
        # Some subjects have separate F-6 and 7-10 documents (e.g., Technologies)
        # Others have all bands in one document (e.g., Geography 7-10)
        all_bands = self.config['bands_f6'] + self.config['bands_7_10']
        
        for doc_path in self.config['curriculum_docs']:
            self.parse_curriculum_document(doc_path, all_bands)
            self.parse_achievement_standards(doc_path, all_bands)
        
        self.generate_as_components()
        
        if self.config.get('glossary_doc'):
            self.parse_glossary(self.config['glossary_doc'])
        
        self.validate()
        
        print("\n" + "=" * 60)
        print("WRITING CSV FILES")
        print("=" * 60)
        
        self.write_curriculum_csvs()
        self.write_achievement_standards_csv()
        self.write_as_components_csv()
        self.write_glossary_csv()
        
        # Phase 3: Post-parse validation
        postcheck = ACARAParsedDataTests(
            curriculum_data=[{
                'code': cd.code, 'description': cd.description, 'elaborations': cd.elaborations
            } for cd in self.content_descriptors],
            achievement_standards=self.achievement_standards,
            as_components=[{'strand': c.strand, 'keywords': c.keywords} for c in self.as_components],
            glossary_terms=self.glossary_terms,
            code_pattern=self.code_pattern
        )
        postcheck.run_all()
        
        print("\n" + "=" * 60)
        print("COMPLETE!")
        print(f"Output: {self.config['output_dir']}")
        print("=" * 60)


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    parser = ACARACurriculumParser(SUBJECT_CONFIG)
    parser.run()
